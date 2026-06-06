# -*- coding: utf-8 -*-
"""Générateur du document d'explication du code (.docx)."""

from __future__ import annotations
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

COL_TITLE = RGBColor(0x1F, 0x3A, 0x5F)
COL_H1 = RGBColor(0x1A, 0x5C, 0x8A)
COL_H2 = RGBColor(0x2E, 0x6B, 0x4F)
COL_H3 = RGBColor(0x8A, 0x55, 0x1A)
COL_CODE = RGBColor(0x22, 0x22, 0x22)
CODE_BG = "F2F3F5"
NOTE_BG = "FFF6E5"


def _shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "6"); el.set(qn("w:color"), "D0D3D7")
        pbdr.append(el)
    pPr.append(pbdr)


def add_runs(p, text):
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            p.add_run(tok)


def H(doc, text, level):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    if level == 0:
        r.font.size = Pt(20); r.font.color.rgb = COL_TITLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        r.font.size = Pt(15); r.font.color.rgb = COL_H1
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    elif level == 2:
        r.font.size = Pt(12.5); r.font.color.rgb = COL_H2
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
    else:
        r.font.size = Pt(11); r.font.color.rgb = COL_H3
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(1)
    return p


def P(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text)
    return p


def BULLET(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text)
    return p


def CODE(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.left_indent = Inches(0.08)
    _shade(p, CODE_BG); _border(p)
    r = p.add_run(code)
    r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = COL_CODE
    return p


def NOTE(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _shade(p, NOTE_BG); _border(p)
    add_runs(p, text)
    return p


# ===========================================================================
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

H(doc, "Données cliniques synthétiques — explication du code", 0)
P(doc, "")

# --------------------------------------------------------------------------
H(doc, "Vue d'ensemble", 1)
P(doc, "Le projet est un **pipeline de génération de données** : il produit un "
       "jeu de patients synthétiques, le valide, l'analyse et l'évalue en "
       "Machine Learning. On génère 6 classes équilibrées (sain, diabète, "
       "dyslipidémie, hypertension, obésité, risque cardiovasculaire).")
P(doc, "L'approche encode des **connaissances médicales** (seuils ADA/ESC/OMS) et "
       "la **structure de dépendance** entre variables, puis échantillonne via "
       "une **copule gaussienne** et un **CTGAN**, avant de **rejeter** tout "
       "enregistrement incohérent.")
P(doc, "Ordre des modules : **config → core → generators → validation → analysis "
       "→ ml_evaluation → visualization → pipeline → run.py**.")

# ===========================================================================
H(doc, "MODULE 1 — config/", 1)
P(doc, "Rôle du module : centraliser toutes les constantes du domaine. Le reste "
       "du code n'écrit aucun chiffre clinique en dur ; il importe depuis config. "
       "Détail ci-dessous fichier par fichier, section par section.")

# ===========================================================================
H(doc, "Fichier 1 — clinical_ranges.py", 1)
P(doc, "But du fichier : décrire **chaque variable clinique** et **chaque seuil** "
       "qui servira à classer, générer et valider les patients. C'est le "
       "référentiel de base. Il est organisé en 8 sections.")

H(doc, "Section 1 — Types (les « moules » de variables)", 2)
P(doc, "Contenu : trois définitions de structure. L'énumération `VariableType` "
       "(continue, binaire, nominale, ordinale) ; la dataclass "
       "`ContinuousVariableSpec` (nom, unité, bornes absolues, bornes normales) ; "
       "la dataclass `CategoricalVariableSpec` (nom, type, liste des modalités).")
P(doc, "Utilité : poser un **format commun**. Toutes les variables seront créées "
       "à partir de ces moules, donc toutes ont les mêmes champs. Le contrôle "
       "`__post_init__` de `ContinuousVariableSpec` vérifie à la création que "
       "`absolute_min ≤ normal_min ≤ normal_max ≤ absolute_max` : une faute de "
       "frappe dans un seuil fait planter le programme immédiatement.")

H(doc, "Section 2 — Variables démographiques", 2)
P(doc, "Contenu : `AGE`, `HEIGHT`, `WEIGHT`, `BMI`, chacune instanciée avec ses "
       "bornes (ex. âge 18–90, IMC normal 18,5–24,9).")
P(doc, "Utilité : fixer le profil de base d'un patient (qui il est physiquement). "
       "Ces bornes servent ensuite au rejet (validation) et au cadrage de la "
       "génération.")

H(doc, "Section 3 — Signes vitaux", 2)
P(doc, "Contenu : `HEART_RATE`, `SBP` (tension systolique), `DBP` (tension "
       "diastolique), `RESP_RATE`, `TEMPERATURE`, avec leurs plages normales et "
       "absolues.")
P(doc, "Utilité : décrire les mesures « au lit du patient ». Notamment SBP/DBP "
       "alimenteront le diagnostic d'hypertension.")

H(doc, "Section 4 — Résultats de laboratoire", 2)
P(doc, "Contenu : `FASTING_GLUCOSE`, `HBA1C`, `TOTAL_CHOL`, `LDL`, `HDL`, "
       "`TRIGLYCERIDES`.")
P(doc, "Utilité : ce sont les variables qui portent les diagnostics de diabète "
       "et de dyslipidémie. Leurs bornes normales définissent la frontière du sain.")

H(doc, "Section 5 — Variables catégorielles et de mode de vie", 2)
P(doc, "Contenu : `SEX` (binaire), `PHYSICAL_ACTIVITY` et `ALCOHOL` et "
       "`DIET_QUALITY` (ordinales), `SMOKING` (nominale), avec la liste de leurs "
       "modalités possibles.")
P(doc, "Utilité : définir les facteurs comportementaux. Ils interviennent surtout "
       "dans le score de risque cardiovasculaire (tabac, sédentarité, etc.).")

H(doc, "Section 6 — Registres exposés", 2)
P(doc, "Contenu : deux dictionnaires, `CONTINUOUS_VARIABLES` et "
       "`CATEGORICAL_VARIABLES`, construits par compréhension `{v.name: v}`.")
P(doc, "Utilité : offrir un **point d'accès unique** au reste du code, qui "
       "récupère une variable par son nom (ex. `CONTINUOUS_VARIABLES[\"bmi\"]`) "
       "sans connaître les détails internes.")

H(doc, "Section 7 — Seuils diagnostiques", 2)
P(doc, "Contenu : les constantes de seuils, nommées par convention `DX_*` "
       "(seuil pathologique strict, ex. `DX_GLUCOSE = 126`) et `RF_*` "
       "(facteur de risque / zone intermédiaire, ex. `RF_GLUCOSE = 100`). Plus "
       "la fonction `hdl_low_threshold(sex)` qui renvoie le seuil de HDL bas "
       "selon le sexe (40 homme / 50 femme).")
P(doc, "Utilité : c'est la **table de décision médicale**, isolée ici pour "
       "n'exister qu'à un seul endroit. Les fichiers `diagnostic_criteria` et "
       "`generation_params` s'y réfèrent.")

H(doc, "Section 8 — Tolérances de cohérence", 2)
P(doc, "Contenu : constantes comme `TOLERANCE_BMI`, `TOLERANCE_FRIEDEWALD`, "
       "`MIN/MAX_PULSE_PRESSURE`.")
P(doc, "Utilité : fixer la marge d'erreur acceptée lors des contrôles de "
       "relations entre variables (module validation) — par ex. vérifier que "
       "l'IMC est bien cohérent avec poids et taille à ±0,5 près.")

# ===========================================================================
H(doc, "Fichier 2 — diagnostic_criteria.py", 1)
P(doc, "But du fichier : répondre par oui/non à « ce patient a-t-il telle "
       "condition ? », à partir des seuils du fichier 1. 5 sections.")

H(doc, "Section 1 — Imports et type d'entrée", 2)
P(doc, "Contenu : import de tous les seuils depuis `clinical_ranges`, et alias "
       "`PatientValues = Mapping[str, Any]` (un patient vu comme un dictionnaire "
       "{nom_variable: valeur}).")
P(doc, "Utilité : garantir qu'aucun seuil n'est redéfini ici (une seule source "
       "de vérité) et définir le format d'entrée commun à tous les prédicats.")

H(doc, "Section 2 — Prédicats diagnostiques stricts", 2)
P(doc, "Contenu : `is_diabetic`, `is_dyslipidemic`, `is_hypertensive`, "
       "`is_obese`. Chacune compare les mesures du patient aux seuils `DX_*`.")
CODE(doc, 'def is_diabetic(p):\n'
          '    return p["fasting_glucose"] >= DX_GLUCOSE or p["hba1c"] >= DX_HBA1C')
P(doc, "Utilité : déterminer l'appartenance aux classes mono-pathologie. "
       "Ces fonctions seront appelées par le module core pour étiqueter chaque "
       "patient généré.")

H(doc, "Section 3 — Facteurs de risque cardiovasculaire", 2)
P(doc, "Contenu : des fonctions privées `_has_glycemic_risk`, `_has_lipid_risk`, "
       "`_has_blood_pressure_risk`, `_has_bmi_risk` (facteurs biologiques, basés "
       "sur les seuils `RF_*`), plus `_has_heart_rate_risk`, `_has_smoking_risk`, "
       "`_has_sedentary_risk`, `_has_alcohol_risk`, `_has_diet_risk` "
       "(comportementaux).")
P(doc, "Utilité : décomposer le risque CV en facteurs élémentaires "
       "indépendants, chacun testable séparément. Le préfixe `_` indique qu'ils "
       "sont internes au fichier.")

H(doc, "Section 4 — count_cv_risk_factors", 2)
P(doc, "Contenu : agrège les facteurs ci-dessus et renvoie le couple "
       "`(nombre_total, nombre_biologiques)`.")
P(doc, "Utilité : fournir le **score** sur lequel repose la règle d'attribution "
       "de la classe risque CV (`total ≥ 3 OU biologiques ≥ 2`), appliquée plus "
       "tard dans core.")

H(doc, "Section 5 — cv_risk_factors_breakdown", 2)
P(doc, "Contenu : renvoie un dictionnaire {nom_du_facteur: True/False}.")
P(doc, "Utilité : version détaillée du score, pour le débogage et les "
       "visualisations (savoir *quels* facteurs précis sont présents).")

# ===========================================================================
H(doc, "Fichier 3 — generation_params.py", 1)
P(doc, "But du fichier : décrire statistiquement **chaque classe de patients** "
       "(quelles moyennes, quelles lois) pour que le générateur sache produire "
       "un diabétique « type », un sain « type », etc. 5 sections.")

H(doc, "Section 1 — Le type ContinuousParam", 2)
P(doc, "Contenu : dataclass `ContinuousParam(distribution, mean, std, "
       "mean_male, mean_female)` et sa méthode `transform(z, sex)`.")
P(doc, "Utilité : représenter la loi d'une variable continue. `transform` "
       "convertit un tirage standard `z ~ N(0,1)` en valeur réelle : `mean+std*z` "
       "pour une normale, formule log-échelle pour une log-normale. Les champs "
       "`mean_male/mean_female` permettent une moyenne dépendante du sexe "
       "(taille, HDL).")

H(doc, "Section 2 — Constantes physiologiques partagées", 2)
P(doc, "Contenu : `HEIGHT_MEAN_MALE/FEMALE`, `HEIGHT_STD`, et les paramètres "
       "`RESP_RATE_PARAM`, `TEMP_PARAM` communs à toutes les classes.")
P(doc, "Utilité : éviter de redéfinir dans chaque classe les variables qui ne "
       "discriminent pas les profils (fréquence respiratoire, température, "
       "modèle de taille).")

H(doc, "Section 3 — Paramètres continus par classe", 2)
P(doc, "Contenu : 6 dictionnaires (`PARAMS_HEALTHY` … `PARAMS_CV_RISK`) donnant "
       "moyenne+écart-type de chaque variable, puis le registre "
       "`CONTINUOUS_PARAMS_BY_CLASS` qui les relie aux classes.")
P(doc, "Utilité : c'est le **cœur du réalisme par classe**. Calibrage clé "
       "(contrainte mono-label) : une classe pathologique ne franchit que *son* "
       "seuil. Ex. un diabétique a une glycémie centrée à 158 mais une tension, "
       "des lipides et un IMC maintenus sous les seuils `RF_*`, pour ne pas "
       "basculer en risque CV. `weight` et `total_chol` sont absents car "
       "**dérivés par formule** (cohérence garantie).")

H(doc, "Section 4 — Paramètres catégoriels par classe", 2)
P(doc, "Contenu : 6 dictionnaires (`PROBS_HEALTHY` …) donnant la probabilité de "
       "chaque modalité (ex. tabac : 45 % de fumeurs en CV_RISK contre 10 % en "
       "sain), puis le registre `CATEGORICAL_PROBS_BY_CLASS`.")
P(doc, "Utilité : reproduire les associations épidémiologiques réelles entre une "
       "classe et les comportements (tabac, alcool, activité, alimentation).")

H(doc, "Section 5 — Auto-validation à l'import", 2)
P(doc, "Contenu : `_validate_categorical_distributions()`, appelée immédiatement "
       "en fin de fichier.")
P(doc, "Utilité : garantir au chargement que chaque distribution catégorielle "
       "somme bien à 1 et n'utilise que des modalités déclarées. Sécurité contre "
       "les erreurs de saisie.")

# ===========================================================================
H(doc, "Fichier 4 — correlations.py", 1)
P(doc, "But du fichier : fournir la **matrice de corrélation** qui dit comment "
       "les variables varient ensemble. Sans elle, les variables seraient "
       "générées indépendamment, ce qui serait irréaliste. 4 sections.")

H(doc, "Section 1 — Variables de la copule", 2)
P(doc, "Contenu : le tuple `COPULA_VARIABLES` (les 13 variables continues "
       "échantillonnées, dans un ordre fixe).")
P(doc, "Utilité : définir l'ordre canonique des lignes/colonnes de la matrice. "
       "`weight` et `total_chol` en sont exclus car dérivés, pas tirés.")

H(doc, "Section 2 — Paires de corrélation", 2)
P(doc, "Contenu : `CORRELATION_PAIRS`, liste de triplets `(var1, var2, ρ)` avec "
       "leur justification physiologique (ex. `bmi↔triglycerides = +0.40`, "
       "`fasting_glucose↔hba1c = +0.85`, `bmi↔hdl = -0.30`).")
P(doc, "Utilité : exprimer les liens connus entre variables. Les paires non "
       "listées sont implicitement de corrélation 0.")

H(doc, "Section 3 — Construction et validation de la matrice", 2)
P(doc, "Contenu : `build_correlation_matrix` (assemble la matrice symétrique à "
       "partir des paires), `is_positive_definite` (teste la factorisation de "
       "Cholesky), `project_to_nearest_correlation` (corrige par projection "
       "spectrale), `ensure_valid_correlation_matrix` (orchestre le tout).")
P(doc, "Utilité : pour échantillonner une normale multivariée, la matrice doit "
       "être **définie positive**. Ces fonctions garantissent que la matrice "
       "fournie l'est — et la réparent sinon (on tronque les valeurs propres "
       "négatives puis on renormalise la diagonale à 1).")

H(doc, "Section 4 — Matrice de référence", 2)
P(doc, "Contenu : `BASE_CORRELATION_MATRIX` (la matrice finale validée) et "
       "`COPULA_INDEX` (dictionnaire nom→indice).")
P(doc, "Utilité : objets prêts à l'emploi consommés directement par le "
       "générateur copule ; `COPULA_INDEX` permet de retrouver la position d'une "
       "variable dans la matrice.")

NOTE(doc, "**À retenir :** config/ = référentiel (clinical_ranges) + règles de "
          "décision (diagnostic_criteria) + paramètres de génération "
          "(generation_params) + structure de corrélation (correlations). Chaque "
          "section a un rôle unique et les valeurs sont validées dès l'import.")

# ===========================================================================
H(doc, "MODULE 2 — core/", 1)
P(doc, "Rôle du module : définir **ce qu'est un patient** (sa structure de "
       "données) et **comment lui attribuer sa classe**. Là où config/ contient "
       "des chiffres, core/ contient les objets manipulés par tout le pipeline. "
       "2 fichiers principaux.")

# ===========================================================================
H(doc, "Fichier 1 — patient_schema.py", 1)
P(doc, "But du fichier : fournir une représentation **typée** d'un "
       "enregistrement patient (les 20 variables + identifiant + classe), avec "
       "les conversions nécessaires pour passer du générateur au DataFrame/CSV. "
       "5 sections.")

H(doc, "Section 1 — Enums (modalités strictes)", 2)
P(doc, "Contenu : `ClassLabel` (les 6 classes) et une énumération par variable "
       "catégorielle — `Sex`, `ActivityLevel`, `SmokingStatus`, "
       "`AlcoholConsumption`, `DietQuality`. Chacune hérite de `str`.")
P(doc, "Utilité : restreindre les valeurs possibles à une liste fermée. On ne "
       "peut pas écrire `smoking = \"fumeur\"` par erreur ; seules les modalités "
       "déclarées sont acceptées. Comme elles héritent de `str`, elles "
       "s'écrivent directement dans un CSV sans conversion.")

H(doc, "Section 2 — La dataclass Patient", 2)
P(doc, "Contenu : la classe `Patient` qui liste tous les champs typés, regroupés "
       "en démographiques, signes vitaux, laboratoire, mode de vie, plus "
       "`patient_id` et `class_label`.")
P(doc, "Utilité : c'est l'**objet central** du projet. Un patient généré devient "
       "une instance de `Patient` ; le typage documente l'unité de chaque champ "
       "(cm, mmHg, mg/dL…) et permet aux outils de détecter les erreurs.")

H(doc, "Section 3 — Conversions", 2)
P(doc, "Contenu : `to_dict()` (convertit en dictionnaire, en remplaçant les "
       "enums par leur valeur texte) et `to_flat_values()` (vue plate des "
       "valeurs cliniques).")
P(doc, "Utilité : faire le pont entre l'objet `Patient` et les autres formats. "
       "`to_dict` sert à construire un DataFrame ou un JSON ; `to_flat_values` "
       "produit l'entrée attendue par les prédicats de `class_assigner` et les "
       "règles de cohérence.")

H(doc, "Section 4 — Construction", 2)
P(doc, "Contenu : `new_id()` (identifiant anonyme aléatoire via `secrets`), "
       "`field_names()` (noms des champs dans l'ordre), `from_mapping(data)` "
       "(construit un `Patient` depuis un dictionnaire en convertissant les "
       "chaînes en enums).")
P(doc, "Utilité : `from_mapping` est le point d'entrée utilisé par les "
       "générateurs : ils produisent des dictionnaires de valeurs, cette méthode "
       "les transforme en objets `Patient` validés (modalité inconnue → "
       "exception). `field_names` garantit un ordre de colonnes stable au CSV.")

H(doc, "Section 5 — Métadonnées", 2)
P(doc, "Contenu : les tuples `CONTINUOUS_FIELDS` et `CATEGORICAL_FIELDS`.")
P(doc, "Utilité : lister explicitement quelles variables sont continues et "
       "lesquelles sont catégorielles. Le reste du pipeline (analyses, "
       "visualisations, ML) s'en sert pour traiter chaque type correctement "
       "sans recoder la liste à chaque fois.")

# ===========================================================================
H(doc, "Fichier 2 — class_assigner.py", 1)
P(doc, "But du fichier : déterminer la classe **unique** d'un patient à partir de "
       "ses valeurs, selon une hiérarchie de priorité. C'est la règle qui "
       "garantit le cadre mono-label. 4 sections.")

H(doc, "Section 1 — Imports et seuils de la règle", 2)
P(doc, "Contenu : import des prédicats de `diagnostic_criteria`, et les "
       "constantes `CV_RISK_MIN_TOTAL_FACTORS = 3` et "
       "`CV_RISK_MIN_BIOLOGICAL_FACTORS = 2`.")
P(doc, "Utilité : rendre explicites et modifiables les deux seuils qui "
       "déclenchent la classe risque cardiovasculaire.")

H(doc, "Section 2 — assign_class (la hiérarchie)", 2)
P(doc, "Contenu : la fonction qui applique 3 étapes dans l'ordre.")
CODE(doc, 'n_total, n_bio = count_cv_risk_factors(p)\n'
          'if n_total >= 3 or n_bio >= 2:   return CV_RISK   # étape 1\n'
          'if is_diabetic(p):               return DIABETES  # étape 2\n'
          'if is_dyslipidemic(p):           return DYSLIPIDEMIA\n'
          'if is_hypertensive(p):           return HYPERTENSION\n'
          'if is_obese(p):                  return OBESITY\n'
          'return HEALTHY                                    # étape 3')
P(doc, "Utilité : l'**ordre est crucial**. On teste d'abord le profil "
       "multifactoriel (risque CV) car il est prioritaire ; sinon les "
       "pathologies uniques dans un ordre fixe ; sinon « sain ». Cela évite "
       "qu'un patient soit éligible à deux classes : il en reçoit toujours "
       "exactement une.")

H(doc, "Section 3 — matches_expected_class", 2)
P(doc, "Contenu : compare la classe inférée par `assign_class` à une classe "
       "attendue, renvoie un booléen.")
P(doc, "Utilité : pilier du **rejection sampling**. Lors de la génération, on "
       "vise une classe cible ; si la classe réellement inférée des valeurs "
       "diffère, le candidat est rejeté. Sert aussi au validateur pour vérifier "
       "la cohérence valeur/étiquette.")

H(doc, "Section 4 — explain_class_assignment", 2)
P(doc, "Contenu : renvoie un dictionnaire détaillé (classe inférée, compteurs de "
       "facteurs, état de chaque prédicat).")
P(doc, "Utilité : outil de **diagnostic/log**, hors boucle critique. Quand un "
       "patient est rejeté, il explique pourquoi (quels critères ont basculé).")

NOTE(doc, "**À retenir :** core/ définit l'objet `Patient` (structure typée + "
          "conversions) et la fonction `assign_class` (hiérarchie de priorité "
          "mono-label). C'est le pont entre les chiffres de config/ et les "
          "patients concrets manipulés par les générateurs et le validateur.")

# ===========================================================================
# MODULE 3 — generators/
# ===========================================================================
H(doc, "MODULE 3 — generators/", 1)
P(doc, "Rôle du module : c'est le module qui **fabrique réellement les "
       "patients**. Il propose deux méthodes de génération différentes mais "
       "interchangeables : la **copule gaussienne** (Méthode 1, purement "
       "statistique) et le **CTGAN** (Méthode 2, un réseau de neurones "
       "génératif). Les deux produisent des candidats puis ne conservent que "
       "ceux qui passent la validation clinique (« rejection sampling »). "
       "Pour éviter de dupliquer du code, une **classe abstraite** commune "
       "(`BaseGenerator`) définit ce qu'un générateur doit savoir faire, et la "
       "boucle qui produit le jeu équilibré complet n'est écrite qu'une seule "
       "fois. Le module contient 3 fichiers.")

# ---------------------------------------------------------------- Fichier 1
H(doc, "Fichier 1 — base_generator.py", 1)
P(doc, "But du fichier : définir le **contrat commun** à tous les générateurs. "
       "C'est une classe abstraite (`ABC`) : elle ne génère rien elle-même, mais "
       "elle impose une structure que copule et CTGAN doivent respecter, et elle "
       "implémente une fois pour toutes la logique partagée. 2 sections.")

H(doc, "Section 1 — La classe abstraite BaseGenerator", 2)
P(doc, "Contenu : une classe héritant de `ABC` qui déclare deux méthodes "
       "**abstraites** (marquées `@abstractmethod`, donc obligatoires à "
       "implémenter) : `fit(training_data)` pour entraîner ou paramétrer le "
       "générateur, et `sample_class(class_label, n)` pour produire `n` patients "
       "valides d'une classe donnée. Elle déclare aussi un attribut "
       "`stats: RejectionStatistics`.")
P(doc, "Utilité : garantir que **n'importe quel générateur s'utilise de la même "
       "façon**. Le reste du pipeline ne sait pas s'il manipule une copule ou un "
       "CTGAN : il appelle `fit` puis `sample_class`. On peut donc ajouter une "
       "3ᵉ méthode demain sans rien changer ailleurs. C'est le principe du "
       "polymorphisme.")

H(doc, "Section 2 — generate_balanced_dataset", 2)
P(doc, "Contenu : une méthode **concrète** (déjà écrite ici) qui boucle sur les "
       "6 classes de `ClassLabel`, appelle `sample_class(classe, m_per_class)` "
       "pour chacune, mesure la durée et le taux de rejet de chaque classe, les "
       "journalise, puis assemble tous les patients en un `DataFrame` dont les "
       "colonnes sont remises dans l'ordre officiel via `Patient.field_names()`.")
CODE(doc, 'for i, class_label in enumerate(ClassLabel, start=1):\n'
          '    patients = self.sample_class(class_label, m_per_class)\n'
          '    all_patients.extend(patients)\n'
          '    # ... log : [i/6] classe ✓ N patients (rejet X%, durée s)\n'
          'df = pd.DataFrame([p.to_dict() for p in all_patients])\n'
          'return df[list(Patient.field_names())]')
P(doc, "Utilité : produire en **un seul appel le jeu équilibré complet** "
       "(6 × M patients, autant par classe, conformément au cahier des charges). "
       "Comme cette logique est identique pour les deux méthodes, elle est "
       "factorisée ici : copule et CTGAN n'ont qu'à fournir `sample_class`, et ils "
       "héritent gratuitement de cette boucle.")

# ---------------------------------------------------------------- Fichier 2
H(doc, "Fichier 2 — gaussian_copula.py (Méthode 1)", 1)
P(doc, "But du fichier : générer des patients par **copule gaussienne**. "
       "L'idée centrale d'une copule est de **séparer deux choses** : d'un côté "
       "la **structure de dépendance** entre variables (la matrice de "
       "corrélation `R` de config/), de l'autre la **forme de chaque variable "
       "prise seule** (sa loi marginale : normale ou log-normale, avec la moyenne "
       "de sa classe). Concrètement : on tire d'abord un vecteur gaussien "
       "corrélé `Z`, puis on transforme chaque composante pour lui donner la "
       "bonne loi clinique, sans casser les corrélations. 5 sections.")

H(doc, "Section 1 — Constante de bruit Friedewald", 2)
P(doc, "Contenu : `FRIEDEWALD_NOISE_STD = 3.0`.")
P(doc, "Utilité : le cholestérol total n'est pas tiré au hasard, il est "
       "**recalculé** par la formule de Friedewald (LDL + HDL + TG/5). On y "
       "ajoute un minuscule bruit aléatoire d'écart-type 3 mg/dL pour que la "
       "valeur ne soit pas parfaitement déterministe — tout en restant sous la "
       "tolérance de validation de ±10 mg/dL (règle R2). C'est un détail de "
       "réalisme.")

H(doc, "Section 2 — _vectorized_marginal_transform", 2)
P(doc, "Contenu : une fonction qui applique la **transformation marginale** à "
       "tout un lot de valeurs d'un coup (en numpy, sans boucle Python). Elle "
       "calcule les moyennes effectives — éventuellement différentes selon le "
       "sexe via `np.where(sexes == \"Male\", ...)` — puis applique `moyenne + "
       "écart-type × z` pour une loi normale, ou la formule log-échelle pour une "
       "log-normale.")
P(doc, "Utilité : c'est l'équivalent vectorisé de `ContinuousParam.transform` vu "
       "dans config/. On transforme des **milliers de candidats simultanément** "
       "au lieu d'un par un : c'est ce qui rend la génération rapide.")

H(doc, "Section 3 — Initialisation (le générateur et Cholesky)", 2)
P(doc, "Contenu : le constructeur `__init__` stocke la matrice de corrélation "
       "(par défaut `BASE_CORRELATION_MATRIX`), crée un générateur aléatoire "
       "reproductible `np.random.default_rng(seed)`, prépare l'objet de "
       "statistiques, et surtout **pré-calcule la factorisation de Cholesky** de "
       "la matrice : `L` tel que `L · Lᵀ = R`.")
CODE(doc, 'self._cholesky_factor = np.linalg.cholesky(self.correlation_matrix)')
P(doc, "Utilité : la factorisation de Cholesky est l'outil mathématique qui "
       "permet de **transformer du bruit indépendant en bruit corrélé**. Si "
       "`Y ~ N(0, I)` (variables indépendantes), alors `Y · Lᵀ` suit `N(0, R)` "
       "(variables corrélées selon R). On la calcule **une seule fois** ici car "
       "elle resservira à chaque lot, ce qui économise beaucoup de calcul.")

H(doc, "Section 4 — sample_class (le rejection sampling)", 2)
P(doc, "Contenu : la méthode centrale. Elle génère des **lots de candidats** "
       "(via `_generate_batch`), puis pour chaque candidat appelle "
       "`validate(candidat, classe)` : si le candidat est valide, on lui attribue "
       "un identifiant, on le convertit en `Patient` et on l'ajoute aux acceptés ; "
       "sinon on enregistre la cause du rejet. La boucle continue jusqu'à obtenir "
       "exactement `n` patients. Un **garde-fou** (`max_total_attempts_factor`, "
       "200 par défaut) lève une erreur si le taux de rejet est anormalement "
       "élevé, pour éviter une boucle infinie.")
CODE(doc, 'result = validate(candidate, class_label)\n'
          'if result.is_valid:\n'
          '    candidate["patient_id"] = Patient.new_id()\n'
          '    accepted.append(Patient.from_mapping(candidate))\n'
          '    self.stats.record_acceptance(class_label)\n'
          'else:\n'
          '    self.stats.record_rejection(class_label, result.failed_rule)')
P(doc, "Utilité : c'est le cœur du **rejection sampling**. On accepte de "
       "« gaspiller » des candidats : on en génère beaucoup, et on ne garde que "
       "ceux qui sont médicalement cohérents ET correctement classés. C'est ce "
       "compromis qui garantit la qualité clinique du jeu final, au prix de "
       "quelques tentatives supplémentaires.")

H(doc, "Section 5 — Méthodes internes et CLI", 2)
P(doc, "Contenu : trois méthodes privées et un point d'entrée ligne de commande.")
BULLET(doc, "`_sample_correlated_z(batch)` : tire `Y ~ N(0, I)` puis renvoie "
            "`Y · Lᵀ`, soit un lot de vecteurs corrélés selon R.")
BULLET(doc, "`_sample_categorical(classe, variable, n)` : tire les variables "
            "qualitatives (sexe, tabac…) selon la distribution multinomiale de la "
            "classe.")
BULLET(doc, "`_generate_batch(classe, taille)` : assemble un lot complet de "
            "candidats. Ordre des étapes : (1) tirer le sexe — car il pilote les "
            "moyennes sexe-dépendantes ; (2) tirer `Z` corrélé ; (3) transformer "
            "chaque variable continue ; (4) tirer les autres catégorielles ; "
            "(5) **dériver** `weight = bmi·(height/100)²` et `total_chol` par "
            "Friedewald ; (6) emballer le tout en dictionnaires.")
P(doc, "Utilité : `_generate_batch` est l'endroit où la cohérence est garantie "
       "**par construction** : poids et cholestérol total sont calculés à partir "
       "des autres variables, donc les règles R1 (Quetelet) et R2 (Friedewald) "
       "sont satisfaites d'office, sans dépendre du hasard. Le `main()` final "
       "permet de lancer la Méthode 1 seule depuis le terminal "
       "(`python -m ...gaussian_copula --m 1000`).")

# ---------------------------------------------------------------- Fichier 3
H(doc, "Fichier 3 — ctgan_generator.py (Méthode 2)", 1)
P(doc, "But du fichier : générer les patients avec un **CTGAN** (Conditional "
       "Tabular GAN), un réseau de neurones génératif spécialisé dans les "
       "données en tableau, fourni par la bibliothèque **SDV**. Contrairement à "
       "la copule (où l'on impose nous-mêmes les lois), le CTGAN **apprend** la "
       "distribution des données à partir d'un jeu d'entraînement — en pratique, "
       "le jeu produit par la copule. Le « Conditional » signifie qu'on peut lui "
       "demander de générer une classe précise. La structure (boucle de rejet, "
       "garde-fou) reste volontairement identique à la copule pour que les deux "
       "méthodes soient comparables. 6 sections.")

H(doc, "Section 1 — Initialisation et graines aléatoires", 2)
P(doc, "Contenu : le constructeur stocke les hyperparamètres d'entraînement "
       "(`epochs`, `batch_size`, `enforce_derived_variables`, seed, garde-fou…) "
       "et prépare l'état interne (`_synthesizer`, `_is_fitted`). La méthode "
       "`_set_seeds` propage la graine à numpy **et à torch** (le moteur de "
       "calcul des réseaux de neurones).")
P(doc, "Utilité : paramétrer l'apprentissage et garantir la **reproductibilité** "
       "malgré le caractère aléatoire d'un réseau de neurones (sans graine fixée, "
       "deux exécutions donneraient des résultats différents).")

H(doc, "Section 2 — _build_metadata", 2)
P(doc, "Contenu : construit l'objet `SingleTableMetadata` de SDV en déclarant "
       "chaque colonne : les variables continues en `numerical`, les "
       "catégorielles et la classe en `categorical`.")
P(doc, "Utilité : le CTGAN a besoin de **savoir le type de chaque colonne** pour "
       "choisir son encodage interne (une variable continue et une variable "
       "catégorielle ne s'apprennent pas de la même façon). On le construit "
       "automatiquement à partir du schéma typé du projet (core/).")

H(doc, "Section 3 — fit (entraînement)", 2)
P(doc, "Contenu : retire la colonne `patient_id` (un identifiant unique n'a "
       "aucune valeur statistique à apprendre), construit la métadonnée, "
       "instancie un `CTGANSynthesizer` avec les hyperparamètres, puis l'entraîne "
       "sur les données via `.fit()`. Marque ensuite `_is_fitted = True`.")
P(doc, "Utilité : c'est la phase d'**apprentissage** : le réseau ajuste ses poids "
       "pour reproduire la distribution jointe des variables du jeu "
       "d'entraînement. Sans cet entraînement préalable, `sample_class` "
       "refuserait de fonctionner.")

H(doc, "Section 4 — Échantillonnage conditionnel et contraintes", 2)
P(doc, "Contenu : `_sample_raw_batch(classe, n)` demande à SDV `n` lignes "
       "**conditionnées** sur une classe (objet `Condition`). "
       "`_enforce_constructive_constraints(df)` **recalcule** ensuite `weight` et "
       "`total_chol` à partir des autres variables (Quetelet et Friedewald).")
P(doc, "Utilité : le conditionnement permet d'obtenir directement la classe "
       "voulue, donc un jeu équilibré. Le recalcul des variables dérivées est "
       "important : le réseau pourrait produire un poids légèrement incohérent "
       "avec l'IMC ; en le **réimposant par formule** (exactement comme la "
       "copule), on garantit R1 et R2 et on rend les deux méthodes réellement "
       "comparables.")

H(doc, "Section 5 — sample_class", 2)
P(doc, "Contenu : la même boucle de rejet que la copule, adaptée au CTGAN : "
       "vérifier que le modèle est entraîné, tirer un lot conditionnel, réimposer "
       "les variables dérivées, **valider** chaque ligne, accepter ou rejeter, "
       "jusqu'à atteindre `n`. Le garde-fou est plus tolérant (facteur 500) et "
       "son message d'erreur suggère d'augmenter `epochs` si le rejet est trop "
       "élevé.")
P(doc, "Utilité : appliquer **exactement le même filtre de validation** aux deux "
       "méthodes. Un avantage caché : si le CTGAN est sous-entraîné, il produit "
       "beaucoup de patients incohérents, le taux de rejet explose et le "
       "garde-fou le signale — c'est un indicateur de qualité du modèle.")

H(doc, "Section 6 — CLI", 2)
P(doc, "Contenu : un `main()` qui charge un CSV d'entraînement (typiquement "
       "celui de la copule), entraîne le CTGAN, génère le jeu et l'enregistre.")
P(doc, "Utilité : pouvoir exécuter la Méthode 2 **isolément**, en aval de la "
       "Méthode 1, sans lancer tout le pipeline.")

NOTE(doc, "**À retenir :** generators/ repose sur une interface commune "
          "(`BaseGenerator`) et deux implémentations. La **copule** impose les "
          "lois et les corrélations, et tire via Cholesky. Le **CTGAN** apprend "
          "la distribution à partir des données de la copule. Les deux passent "
          "par le **même rejection sampling** (générer → valider → ne garder que "
          "le cohérent) et réimposent par formule le poids et le cholestérol "
          "total, ce qui les rend rigoureusement comparables.")

# ===========================================================================
# MODULE 4 — validation/
# ===========================================================================
H(doc, "MODULE 4 — validation/", 1)
P(doc, "Rôle du module : c'est le **filtre de qualité clinique**. À chaque fois "
       "qu'un générateur produit un patient candidat, il le soumet à ce module, "
       "qui répond « valide » ou « rejeté ». La stratégie n'est jamais de "
       "**corriger** un patient incohérent, mais de le **rejeter** purement et "
       "simplement (le générateur en produira un autre). Les contrôles sont "
       "organisés en **cascade** : on les enchaîne dans un ordre précis et on "
       "s'arrête au premier échec. Le module mesure aussi le coût de ce rejet "
       "(combien de candidats jetés, pour quelle raison). 5 fichiers.")

# ---------------------------------------------------------------- Fichier 1
H(doc, "Fichier 1 — validator.py (le point d'entrée)", 1)
P(doc, "But du fichier : offrir **une seule fonction** `validate()` qui enchaîne "
       "tous les contrôles des autres fichiers. C'est ce que les générateurs "
       "appellent. 2 sections.")

H(doc, "Section 1 — ValidationResult", 2)
P(doc, "Contenu : une dataclass immuable qui porte trois informations — "
       "`is_valid` (booléen), `failed_rule` (le nom de la première règle violée, "
       "ou None) et `detail` (précision éventuelle, ex. la variable hors borne). "
       "Deux fabriques pratiques l'accompagnent : `ok()` et `failed(rule, detail)`.")
P(doc, "Utilité : renvoyer un résultat **structuré** plutôt qu'un simple "
       "vrai/faux. Savoir *quelle* règle a rejeté un patient permet de remplir "
       "les statistiques de rejet (quel contrôle élimine le plus de candidats) et "
       "de déboguer une classe mal calibrée.")

H(doc, "Section 2 — validate (la cascade)", 2)
P(doc, "Contenu : la fonction qui applique les contrôles dans l'ordre et retourne "
       "au **premier échec** : (1) bornes physiologiques absolues, (2) les règles "
       "inter-variables R1 → R4, (3) la cohérence valeurs/classe.")
CODE(doc, 'bounds_fault = check_absolute_bounds(p)\n'
          'if bounds_fault is not None:\n'
          '    return ValidationResult.failed(RULE_BOUNDS, detail=bounds_fault)\n'
          'for rule_name, rule_fn in INTER_VARIABLE_RULES:\n'
          '    if not rule_fn(p):\n'
          '        return ValidationResult.failed(rule_name)\n'
          'if not check_class_coherence(p, expected_class):\n'
          '    return ValidationResult.failed(RULE_CLASS)\n'
          'return ValidationResult.ok()')
P(doc, "Utilité : l'**ordre est délibérément optimisé pour la performance**. On "
       "commence par les bornes (test le moins coûteux, qui élimine vite les "
       "candidats aberrants), puis les règles R1–R4 (cohérences locales peu "
       "coûteuses), et on finit par la cohérence de classe (la plus coûteuse, car "
       "elle réexécute tous les prédicats diagnostiques). Comme on s'arrête au "
       "premier échec, on évite des calculs inutiles — ce qui compte beaucoup "
       "quand on valide des millions de candidats.")

# ---------------------------------------------------------------- Fichier 2
H(doc, "Fichier 2 — physiological_bounds.py", 1)
P(doc, "But du fichier : première étape de la cascade — vérifier que chaque "
       "valeur reste dans les **limites physiologiques absolues** (et que les "
       "catégories existent). Les bornes viennent toutes de config/, aucune n'est "
       "redéfinie ici. 3 sections.")

H(doc, "Section 1 — violated_continuous_bound", 2)
P(doc, "Contenu : parcourt toutes les variables continues et renvoie le **nom de "
       "la première** qui sort de `[absolute_min, absolute_max]`. Une valeur "
       "manquante, `NaN` ou infinie est aussi traitée comme une violation.")
P(doc, "Utilité : éliminer immédiatement les patients physiquement impossibles "
       "(âge négatif, glycémie à 600, valeur non numérique…). Renvoyer le nom de "
       "la variable fautive aide au diagnostic.")

H(doc, "Section 2 — violated_categorical_modality", 2)
P(doc, "Contenu : vérifie que chaque variable catégorielle prend bien une "
       "modalité **déclarée** dans config/ (pas de valeur inconnue).")
P(doc, "Utilité : se prémunir contre un générateur qui produirait une modalité "
       "inattendue (ex. un statut tabagique mal orthographié).")

H(doc, "Section 3 — check_absolute_bounds", 2)
P(doc, "Contenu : combine les deux contrôles précédents et, en cas d'échec, "
       "renvoie le nom de la variable **préfixé** par son type "
       "(`continuous:` ou `categorical:`).")
P(doc, "Utilité : fournir au validateur un **seul appel** pour toute la première "
       "étape, avec une cause lisible.")

# ---------------------------------------------------------------- Fichier 3
H(doc, "Fichier 3 — inter_variable_rules.py (règles R1 → R4)", 1)
P(doc, "But du fichier : deuxième étape — vérifier la **cohérence entre "
       "variables**. Un patient peut avoir chaque valeur dans les bornes mais des "
       "valeurs incohérentes entre elles (ex. un IMC qui ne correspond pas à sa "
       "taille et son poids). Chaque règle est une fonction qui renvoie True si "
       "la cohérence est respectée. 5 sections.")

H(doc, "Section 1 — R1, Quetelet (IMC)", 2)
P(doc, "Contenu : `check_quetelet` recalcule l'IMC attendu = poids / taille² et "
       "vérifie qu'il colle à l'IMC enregistré à ±0,5 près.")
P(doc, "Utilité : garantir la cohérence **IMC ↔ taille/poids** — exactement "
       "l'exemple cité dans le cahier des charges (« l'IMC doit correspondre à "
       "la taille et au poids »).")

H(doc, "Section 2 — R2, Friedewald (lipides)", 2)
P(doc, "Contenu : `check_friedewald` vérifie que le cholestérol total ≈ LDL + HDL "
       "+ TG/5 (±10 mg/dL). **Exception** : si les triglycérides ≥ 400, la "
       "formule n'est plus valable médicalement, et la règle est considérée comme "
       "passante.")
P(doc, "Utilité : garantir la cohérence du **bilan lipidique**. L'exception "
       "reflète une limite réelle de la formule de Friedewald, ce qui montre la "
       "rigueur médicale du contrôle.")

H(doc, "Section 3 — R3, Pression différentielle", 2)
P(doc, "Contenu : `check_pulse_pressure` vérifie que l'écart PAS − PAD est compris "
       "entre 20 et 100 mmHg.")
P(doc, "Utilité : la pression différentielle reflète l'élasticité des artères ; "
       "une valeur hors de cet intervalle est incompatible avec un patient stable. "
       "On écarte ainsi des combinaisons systolique/diastolique aberrantes.")

H(doc, "Section 4 — R4, Cohérence glycémie / HbA1c", 2)
P(doc, "Contenu : `check_glucose_hba1c_coherence` vérifie trois implications "
       "simultanément. En résumé : une HbA1c (marqueur **chronique**, ~3 mois) et "
       "une glycémie ponctuelle ne doivent pas se contredire — par exemple une "
       "HbA1c diabétique avec une glycémie très basse, ou l'inverse.")
CODE(doc, '# une implication "A ⇒ B" est violée si (A et non B)\n'
          'if hba1c >= DX_HBA1C and glucose < 100:  return False\n'
          'if hba1c <  RF_HBA1C and glucose >= 140: return False\n'
          'if glucose >= DX_GLUCOSE and hba1c < RF_HBA1C: return False')
P(doc, "Utilité : éviter les profils diabétiques médicalement impossibles. C'est "
       "la règle la plus subtile car elle relie deux marqueurs de la même "
       "maladie mesurés à des échelles de temps différentes.")

H(doc, "Section 5 — Registre des règles", 2)
P(doc, "Contenu : `INTER_VARIABLE_RULES`, un tuple ordonné de couples "
       "`(nom, fonction)` listant R1, R2, R3, R4.")
P(doc, "Utilité : permettre au validateur de **parcourir les règles en boucle** "
       "dans l'ordre, sans les appeler une par une en dur. Ajouter une règle R5 "
       "se ferait en une ligne ici.")

# ---------------------------------------------------------------- Fichier 4
H(doc, "Fichier 4 — class_coherence.py", 1)
P(doc, "But du fichier : troisième et dernière étape — vérifier que la **classe "
       "annoncée correspond aux valeurs**. 1 section.")

H(doc, "Section 1 — check_class_coherence", 2)
P(doc, "Contenu : applique `assign_class(p)` (la hiérarchie de core/) sur les "
       "valeurs du patient et vérifie que le résultat est **strictement égal** à "
       "la classe attendue.")
CODE(doc, 'return assign_class(p) == expected_class')
P(doc, "Utilité : c'est la **pierre angulaire du cadre mono-label**. Elle "
       "garantit que l'étiquette stockée dans le dataset est exactement celle "
       "qu'un clinicien (ou un classifieur déterministe) déduirait des valeurs. "
       "C'est aussi ce contrôle qui rejette un candidat généré pour la classe "
       "« diabète » mais dont les valeurs, au final, ressembleraient davantage à "
       "un « risque cardiovasculaire ».")

# ---------------------------------------------------------------- Fichier 5
H(doc, "Fichier 5 — statistics.py", 1)
P(doc, "But du fichier : **mesurer** la génération — combien de candidats "
       "tentés, acceptés, rejetés, et pour quelle raison. 1 section principale.")

H(doc, "Section 1 — RejectionStatistics", 2)
P(doc, "Contenu : une dataclass de compteurs (`Counter`) : tentatives par classe, "
       "acceptés par classe, rejets par règle, et rejets croisés par classe × "
       "règle. Des méthodes de mise à jour (`record_attempt`, `record_acceptance`, "
       "`record_rejection`) sont appelées dans la boucle de génération, et des "
       "méthodes de lecture calculent les indicateurs : `global_rejection_rate`, "
       "`rejection_rate_by_class`, `average_attempts_per_accepted`, et `to_report` "
       "(synthèse sérialisable en JSON).")
P(doc, "Utilité : deux usages. D'abord le **pilotage** : un taux de rejet qui "
       "explose signale une classe mal paramétrée ou un CTGAN sous-entraîné. "
       "Ensuite la **documentation scientifique** : le rapport final peut montrer "
       "l'efficacité de la génération (taux de rejet global, règle la plus "
       "limitante, nombre moyen de tentatives par patient accepté).")

NOTE(doc, "**À retenir :** validation/ est une **cascade** orchestrée par "
          "`validate()` : bornes physiologiques → cohérences R1–R4 → cohérence de "
          "classe, arrêt au premier échec (du moins coûteux au plus coûteux). "
          "`class_coherence` garantit le mono-label, et `statistics` mesure tout "
          "le processus. C'est ce module qui assure la **plausibilité clinique** "
          "exigée par le cahier des charges.")

# ===========================================================================
# MODULE 5 — analysis/
# ===========================================================================
H(doc, "MODULE 5 — analysis/", 1)
P(doc, "Rôle du module : une fois le jeu généré et validé, il faut **prouver "
       "statistiquement qu'il est de bonne qualité**. Ce module ne génère ni ne "
       "rejette rien : il **analyse**. Il répond à trois questions du cahier des "
       "charges : (1) à quoi ressemblent les distributions des variables ? "
       "(2) les corrélations attendues sont-elles bien présentes ? (3) les "
       "classes correspondent-elles aux connaissances épidémiologiques ? Un "
       "quatrième fichier compare les deux méthodes de génération entre elles. "
       "Toutes les fonctions produisent des résultats **sérialisables en JSON** "
       "pour alimenter le rapport scientifique final. 4 fichiers.")

# ---------------------------------------------------------------- Fichier 1
H(doc, "Fichier 1 — descriptive_stats.py", 1)
P(doc, "But du fichier : calculer les **statistiques descriptives** du jeu — les "
       "indicateurs classiques (moyenne, médiane, écart-type, quartiles) qui "
       "résument chaque variable, globalement et par classe. 3 sections.")

H(doc, "Section 1 — Statistiques sur les variables continues", 2)
P(doc, "Contenu : `describe_continuous_global(df)` calcule pour chaque variable "
       "continue : effectif, moyenne, écart-type, min, quartiles (q25, médiane, "
       "q75), max et l'IQR (écart interquartile = q75 − q25). "
       "`describe_continuous_by_class(df)` produit les mêmes indicateurs mais "
       "**séparés par classe** (résultat indexé par couple classe × variable).")
P(doc, "Utilité : décrire la tendance centrale et la dispersion de chaque mesure. "
       "La version par classe est la plus parlante : elle montre par exemple que "
       "la glycémie moyenne des diabétiques est bien plus haute que celle des "
       "sains. C'est la base du tableau « caractéristiques de l'échantillon ».")

H(doc, "Section 2 — Statistiques sur les variables catégorielles", 2)
P(doc, "Contenu : `describe_categorical_global(df)` renvoie, pour chaque variable "
       "qualitative, la **proportion** de chaque modalité. "
       "`describe_categorical_by_class(df)` fait de même par classe, à l'aide d'un "
       "tableau croisé (`pd.crosstab`) normalisé par ligne.")
P(doc, "Utilité : décrire la répartition des facteurs de mode de vie (tabac, "
       "alcool, activité…) et vérifier qu'elle varie comme attendu selon la "
       "classe (ex. davantage de fumeurs chez les patients à risque CV).")

H(doc, "Section 3 — build_descriptive_report", 2)
P(doc, "Contenu : agrège tout ce qui précède en un **dictionnaire structuré** : "
       "effectif total, effectifs par classe, stats continues globales et par "
       "classe, proportions catégorielles globales et par classe. Les structures "
       "à index multiple sont **aplaties** en dictionnaires imbriqués pour rester "
       "compatibles JSON.")
P(doc, "Utilité : produire un rapport unique, exportable et réutilisable "
       "(article, reproductibilité). L'aplatissement explicite évite l'erreur "
       "classique : JSON n'accepte pas les clés sous forme de tuple.")

# ---------------------------------------------------------------- Fichier 2
H(doc, "Fichier 2 — correlation_analysis.py", 1)
P(doc, "But du fichier : analyser les **corrélations réellement présentes** dans "
       "le jeu généré, et les comparer soit à la cible imposée à la copule, soit "
       "à une autre méthode. 3 sections.")

H(doc, "Section 1 — Matrices de corrélation", 2)
P(doc, "Contenu : `pearson_correlation_matrix` (corrélations **linéaires**, la "
       "référence classique), `spearman_correlation_matrix` (corrélations "
       "**monotones**, robustes aux valeurs extrêmes et aux distributions "
       "asymétriques comme les triglycérides), et `correlation_matrix_by_class` "
       "(une matrice par classe).")
P(doc, "Utilité : mesurer comment les variables varient ensemble dans les "
       "données effectivement produites. Avoir les deux mesures (Pearson et "
       "Spearman) est rigoureux : si elles divergent fortement, c'est le signe "
       "d'une relation non linéaire. La version par classe révèle des liens qui "
       "ne s'expriment que dans certains profils (ex. glycémie ↔ HbA1c chez les "
       "diabétiques).")

H(doc, "Section 2 — Comparaison entre matrices", 2)
P(doc, "Contenu : `correlation_difference(A, B)` calcule l'écart **terme à "
       "terme** (A − B) entre deux matrices ; `frobenius_distance(A, B)` en donne "
       "un **résumé en un seul nombre** (la norme de Frobenius, c'est-à-dire la "
       "racine de la somme des carrés des écarts).")
P(doc, "Utilité : quantifier à quel point deux structures de dépendance se "
       "ressemblent. Plus la distance de Frobenius est proche de 0, plus les deux "
       "jeux partagent les mêmes corrélations. C'est l'indicateur global utilisé "
       "pour juger si le CTGAN a bien appris la structure de la copule.")

H(doc, "Section 3 — empirical_vs_target_correlation", 2)
P(doc, "Contenu : compare, **paire par paire**, la corrélation observée dans le "
       "jeu à celle qui avait été **imposée** à la copule "
       "(`BASE_CORRELATION_MATRIX`). Renvoie un tableau avec la valeur cible, la "
       "valeur empirique et leur écart (delta), trié par |delta| décroissant.")
P(doc, "Utilité : **vérification directe** que la copule a bien fait son travail. "
       "Si l'on demandait `bmi ↔ triglycerides = 0,40` et qu'on retrouve ~0,40 "
       "dans les données, la méthode est validée. Le tri par écart met en haut "
       "les paires les plus éloignées de la cible, à examiner en priorité.")

# ---------------------------------------------------------------- Fichier 3
H(doc, "Fichier 3 — epidemiological_validation.py", 1)
P(doc, "But du fichier : vérifier que les classes générées respectent les "
       "**patterns cliniques connus** — autrement dit, confronter les données "
       "non pas à leurs propres règles internes, mais au **savoir médical "
       "externe**. 3 sections.")

H(doc, "Section 1 — Patterns attendus", 2)
P(doc, "Contenu : deux structures de référence. `EXPECTED_PATTERNS` associe à "
       "chaque classe des attentes sur les moyennes (ex. chez les diabétiques, "
       "glycémie moyenne ≥ 126 ; chez les sains, IMC moyen ≤ 25). "
       "`EXPECTED_AGE_ORDERING` liste des comparaisons d'âge attendues (les "
       "classes pathologiques chroniques sont en moyenne plus âgées que les "
       "sains).")
P(doc, "Utilité : encoder, sous forme de données, ce qu'un épidémiologiste "
       "attendrait du jeu. Cela rend la validation **objective et "
       "automatisable**.")

H(doc, "Section 2 — Fonctions de vérification", 2)
P(doc, "Contenu : `check_expected_pattern` teste une attente précise (la moyenne "
       "d'une variable dans une classe est-elle ≥ ou ≤ un seuil ?). "
       "`validate_all_patterns` les lance toutes et renvoie un tableau de "
       "résultats. `check_age_ordering` vérifie les comparaisons d'âge.")
P(doc, "Utilité : produire un **bilan de conformité** clair (chaque attente "
       "passe ou échoue), directement présentable dans le rapport.")

H(doc, "Section 3 — build_epidemiological_report", 2)
P(doc, "Contenu : synthèse — la liste détaillée des patterns, le nombre validés "
       "sur le total, et le résultat de l'ordre des âges.")
P(doc, "Utilité : fournir une **preuve chiffrée du réalisme clinique** "
       "(« X patterns sur Y validés »), un argument fort pour l'article.")

# ---------------------------------------------------------------- Fichier 4
H(doc, "Fichier 4 — method_comparison.py", 1)
P(doc, "But du fichier : comparer directement **la copule (Méthode 1) et le "
       "CTGAN (Méthode 2)** sur trois axes complémentaires. 4 sections.")

H(doc, "Section 1 — Comparaison des marginales", 2)
P(doc, "Contenu : `compare_continuous_means` met côte à côte, pour chaque classe "
       "et chaque variable continue, les moyennes et écarts-types des deux "
       "méthodes, et calcule l'écart absolu et l'écart relatif.")
P(doc, "Utilité : voir si le CTGAN reproduit les **niveaux** de la copule "
       "(les diabétiques du CTGAN ont-ils la même glycémie moyenne ?).")

H(doc, "Section 2 — Comparaison des distributions catégorielles", 2)
P(doc, "Contenu : `compare_categorical_distributions` confronte, pour chaque "
       "variable qualitative, les proportions de chaque modalité par classe entre "
       "les deux jeux (avec le delta).")
P(doc, "Utilité : vérifier que les fréquences comportementales (tabac, alcool…) "
       "sont préservées d'une méthode à l'autre.")

H(doc, "Section 3 — Comparaison de la structure de corrélation", 2)
P(doc, "Contenu : `compare_correlation_structures` calcule les deux matrices "
       "Pearson, leur différence, leur distance de Frobenius, et extrait le "
       "**top 10 des paires** qui divergent le plus.")
P(doc, "Utilité : localiser précisément **où** les deux méthodes diffèrent dans "
       "les relations entre variables.")

H(doc, "Section 4 — build_method_comparison_report", 2)
P(doc, "Contenu : synthèse JSON regroupant effectifs, écarts de moyennes, "
       "distance de Frobenius et top divergences (avec aplatissement des index "
       "multiples).")
P(doc, "Utilité : alimenter la **section comparative** de l'article et conclure "
       "sur la fidélité du CTGAN vis-à-vis de la copule.")

NOTE(doc, "**À retenir :** analysis/ apporte la **preuve statistique** de la "
          "qualité du jeu : il décrit les variables (descriptive_stats), vérifie "
          "les corrélations voulues (correlation_analysis), confronte les classes "
          "au savoir clinique externe (epidemiological_validation) et compare les "
          "deux méthodes (method_comparison). Tous les résultats sont exportés en "
          "JSON pour le rapport.")

# ===========================================================================
# MODULE 6 — ml_evaluation/
# ===========================================================================
H(doc, "MODULE 6 — ml_evaluation/", 1)
P(doc, "Rôle du module : répondre à la question finale du cahier des charges — "
       "**le jeu synthétique est-il vraiment utile pour entraîner une IA ?** On "
       "entraîne donc des modèles de classification à **prédire la classe d'un "
       "patient à partir de ses variables**. Si les modèles y parviennent bien, "
       "c'est que les données portent un signal cohérent et exploitable. Le "
       "module suit trois étapes : **préparer** les données pour les modèles, "
       "**définir** les modèles, puis les **entraîner et évaluer**. 3 fichiers.")

# ---------------------------------------------------------------- Fichier 1
H(doc, "Fichier 1 — preprocessing.py", 1)
P(doc, "But du fichier : transformer les patients en **matrice numérique** "
       "exploitable par les algorithmes. Les modèles ne comprennent que des "
       "nombres : il faut donc encoder les variables catégorielles et mettre les "
       "variables continues à la même échelle. 2 sections.")

H(doc, "Section 1 — Catégorisation des features", 2)
P(doc, "Contenu : des listes qui classent les variables par type de traitement. "
       "`NOMINAL_FIELDS` = sexe, tabac (sans ordre) ; `ORDINAL_FIELDS` = activité, "
       "alcool, qualité de l'alimentation, avec `ORDINAL_CATEGORIES` qui fixe "
       "leur ordre (ex. Sedentary < Moderate < High) ; `EXCLUDED_FROM_FEATURES` = "
       "l'identifiant et la classe, qui ne doivent pas servir de prédicteurs.")
P(doc, "Utilité : appliquer le **bon encodage à chaque type**. Une variable "
       "ordinale doit conserver son ordre (on l'encode 0, 1, 2) ; une variable "
       "nominale n'a pas d'ordre, on la transformera en colonnes binaires. "
       "Exclure `class_label` des features est essentiel : c'est ce qu'on veut "
       "prédire, pas ce qu'on donne en entrée.")

H(doc, "Section 2 — Le préprocesseur et la séparation X/y", 2)
P(doc, "Contenu : `build_preprocessor()` retourne un `ColumnTransformer` sklearn "
       "qui applique en parallèle : `StandardScaler` sur les continues "
       "(centrage-réduction), `OneHotEncoder` sur les nominales (une colonne "
       "binaire par modalité), `OrdinalEncoder` sur les ordinales (un entier par "
       "rang). `prepare_features_and_target(df)` sépare le DataFrame en features "
       "`X` et cible `y` (la classe).")
CODE(doc, 'ColumnTransformer([\n'
          '    ("continuous", StandardScaler(), CONTINUOUS_FIELDS),\n'
          '    ("nominal",    OneHotEncoder(drop="first"), NOMINAL_FIELDS),\n'
          '    ("ordinal",    OrdinalEncoder(categories=...), ORDINAL_FIELDS),\n'
          '])')
P(doc, "Utilité : standardiser les continues évite qu'une variable à grande "
       "échelle (ex. glycémie ~100) écrase une variable à petite échelle (ex. "
       "HbA1c ~5) dans les modèles sensibles à l'échelle. Surtout, ce "
       "préprocesseur sera **intégré dans le pipeline du modèle** (fichier "
       "suivant), ce qui évite la **fuite de données** : les paramètres de "
       "normalisation sont calculés sur le train uniquement, jamais sur le test.")

# ---------------------------------------------------------------- Fichier 2
H(doc, "Fichier 2 — models.py", 1)
P(doc, "But du fichier : définir les **trois familles de modèles** demandées par "
       "le cahier des charges, chacune représentant une approche différente de "
       "l'apprentissage. 2 sections.")

H(doc, "Section 1 — Les fabriques de modèles", 2)
P(doc, "Contenu : trois fonctions, chacune retournant un `Pipeline` sklearn "
       "(préprocesseur + classifieur).")
BULLET(doc, "`make_logistic_regression` : **régression logistique** "
            "multinomiale — un modèle linéaire simple et interprétable, qui sert "
            "de référence (baseline).")
BULLET(doc, "`make_random_forest` : **forêt aléatoire** de 200 arbres — capture "
            "les relations non linéaires et les interactions entre variables sans "
            "réglage lourd.")
BULLET(doc, "`make_mlp` : **réseau de neurones** (perceptron multicouche, 2 "
            "couches cachées de 64 et 32 neurones) avec régularisation L2.")
P(doc, "Utilité : couvrir tout le **spectre demandé** (linéaire, ensembliste, "
       "neuronal) pour comparer leurs performances. Le fait d'emballer chacun "
       "dans un `Pipeline` incluant le préprocesseur garantit que la "
       "normalisation et l'encodage sont refaits proprement à chaque "
       "entraînement, sans fuite. (Un commentaire explique aussi pourquoi "
       "l'early-stopping du MLP est désactivé : un bug sklearn avec des cibles "
       "sous forme de chaînes ; la régularisation L2 compense.)")

H(doc, "Section 2 — Les registres", 2)
P(doc, "Contenu : `MODEL_FACTORIES` (nom interne → fonction de fabrication) et "
       "`MODEL_DISPLAY_NAMES` (nom interne → libellé lisible en français).")
P(doc, "Utilité : permettre au reste du code (évaluation, figures) d'**itérer "
       "automatiquement sur tous les modèles** sans les nommer un par un. Ajouter "
       "un 4ᵉ modèle ne demanderait qu'une ligne ici.")

# ---------------------------------------------------------------- Fichier 3
H(doc, "Fichier 3 — evaluation.py", 1)
P(doc, "But du fichier : **entraîner et mesurer** les modèles selon trois "
       "protocoles de rigueur croissante, et produire un rapport de métriques. "
       "4 sections.")

H(doc, "Section 1 — evaluate_model (la fonction de mesure)", 2)
P(doc, "Contenu : entraîne un modèle sur (X_train, y_train), prédit sur le test, "
       "et calcule un ensemble de métriques : **accuracy** (taux de bonnes "
       "réponses), **F1 macro** (moyenne équilibrée précision/rappel sur les "
       "classes), **F1 par classe**, **ROC AUC** (qualité du classement "
       "probabiliste) et la **matrice de confusion**.")
P(doc, "Utilité : c'est la brique de mesure réutilisée par tous les protocoles. "
       "Le F1 macro est particulièrement adapté ici : comme le jeu est équilibré, "
       "il donne le même poids à chaque classe. La gestion d'erreur sur l'AUC "
       "(renvoyer NaN si une classe manque) évite de faire planter toute "
       "l'évaluation pour un cas limite.")

H(doc, "Section 2 — Protocole 1 : split train/test unique", 2)
P(doc, "Contenu : `evaluate_all_models_on_split` découpe le jeu en 80 % "
       "entraînement / 20 % test de façon **stratifiée** (mêmes proportions de "
       "classes des deux côtés), puis évalue les 3 modèles.")
P(doc, "Utilité : la mesure de performance standard, rapide. La stratification "
       "garantit que chaque classe est représentée dans le test.")

H(doc, "Section 3 — Protocole 2 : validation croisée", 2)
P(doc, "Contenu : `cross_validate_models` applique une **validation croisée "
       "5-fold stratifiée** : le jeu est découpé en 5 parts, chaque modèle est "
       "entraîné 5 fois (4 parts pour entraîner, 1 pour tester, en tournant), et "
       "on rapporte la moyenne et l'écart-type du score.")
P(doc, "Utilité : vérifier que la performance est **stable** et ne dépend pas "
       "d'un découpage chanceux. Un faible écart-type entre les 5 folds est un "
       "gage de fiabilité.")

H(doc, "Section 4 — Protocole 3 : évaluation croisée entre méthodes + rapport", 2)
P(doc, "Contenu : `cross_method_evaluation` entraîne sur un jeu (ex. la copule) "
       "et teste sur l'autre (le CTGAN), et inversement. "
       "`build_ml_evaluation_report` regroupe les trois protocoles en un rapport "
       "JSON.")
P(doc, "Utilité : mesurer la **transférabilité** entre les deux méthodes — un "
       "modèle appris sur les données de la copule sait-il classer les données du "
       "CTGAN ? Une bonne transférabilité prouve que les deux méthodes ont "
       "capturé le même signal clinique sous-jacent, ce qui renforce la "
       "crédibilité du jeu.")

NOTE(doc, "**À retenir :** ml_evaluation/ **encode** les données (preprocessing, "
          "avec un préprocesseur intégré au pipeline pour éviter la fuite), "
          "**définit** trois modèles complémentaires (models : logistique, forêt, "
          "réseau) et les **évalue** selon trois protocoles de rigueur croissante "
          "(evaluation : split, validation croisée, croisé entre méthodes). De "
          "bonnes performances démontrent que le jeu synthétique est exploitable "
          "pour l'IA médicale.")

doc.save("explication_code.docx")
print("OK -> explication_code.docx")
